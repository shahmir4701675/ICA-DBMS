"""
generate_report.py
Generates a professional DBMS Final Project Report (.docx) for
Intelligent Controls & Automation (ICA) — Flask + Supabase project.

Run:
    pip install python-docx
    python generate_report.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def set_font(run, name="Arial", size=12, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_page_number(doc):
    """Add 'Page X of Y' to every page footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    run = para.add_run("Page ")
    set_font(run, size=10, color=(100, 100, 100))

    # PAGE field
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run2 = para.add_run()
    run2._r.append(fld_char)
    run2._r.append(instr)
    run2._r.append(fld_char2)
    set_font(run2, size=10, color=(100, 100, 100))

    run3 = para.add_run(" of ")
    set_font(run3, size=10, color=(100, 100, 100))

    # NUMPAGES field
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.text = " NUMPAGES "
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
    run4 = para.add_run()
    run4._r.append(fc); run4._r.append(it); run4._r.append(fc2)
    set_font(run4, size=10, color=(100, 100, 100))


def add_heading(doc, text, level=1):
    """Add a styled numbered heading."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(31, 78, 121)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(21, 96, 130)
        else:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(50, 50, 50)
    return p


def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=12, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=12)
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_schema_table(doc, headers, rows):
    """Add a styled table for schema display."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.name = "Arial"
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Blue header shading
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            row_cells[c_idx].paragraphs[0].runs[0].font.name = "Arial"
            row_cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
            # Alternate row shading
            if r_idx % 2 == 1:
                tc = row_cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D6E4F0")
                tcPr.append(shd)

    doc.add_paragraph()
    return table


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 150)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Light gray background effect via paragraph border
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build_report():
    doc = Document()

    # ── Page Setup ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    # ── Default paragraph style ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)

    add_page_number(doc)

    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(
        "Database Management System\nFinal Project Report"
    )
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        "Intelligent Controls & Automation (ICA)\n"
        "Service Management Information System"
    )
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(16)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(21, 96, 130)

    doc.add_paragraph()
    doc.add_paragraph()

    meta_lines = [
        ("Course:", "Database Management Systems (DBMS)"),
        ("Semester:", "4th Semester — BS Computer Science"),
        ("Technology Stack:", "PostgreSQL (Supabase) · Flask · React.js"),
        ("Date:", "May 2026"),
    ]
    for label, value in meta_lines:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = mp.add_run(f"{label}  ")
        lr.font.name = "Arial"; lr.font.size = Pt(12); lr.font.bold = True
        vr = mp.add_run(value)
        vr.font.name = "Arial"; vr.font.size = Pt(12)

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS (manual)
    # =========================================================================
    add_heading(doc, "Table of Contents", level=1)

    toc_entries = [
        ("1.", "Introduction", "3"),
        ("2.", "System Overview", "4"),
        ("3.", "Database Design", "5"),
        ("3.1", "Tables & Schema", "5"),
        ("3.2", "Constraints", "8"),
        ("4.", "Data Manipulation Operations", "9"),
        ("5.", "Query Implementation", "11"),
        ("6.", "Frontend & Database Connectivity", "14"),
        ("7.", "Reports Generated", "15"),
        ("8.", "User Roles & Permissions", "16"),
        ("9.", "Testing", "17"),
        ("10.", "Conclusion", "18"),
        ("11.", "References", "19"),
    ]
    for num, title, page in toc_entries:
        tp = doc.add_paragraph()
        tp.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        run_n = tp.add_run(f"{num}  {title}")
        run_n.font.name = "Arial"; run_n.font.size = Pt(12)
        run_dots = tp.add_run("\t" + page)
        run_dots.font.name = "Arial"; run_dots.font.size = Pt(12)
        run_dots.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # =========================================================================
    # 1. INTRODUCTION
    # =========================================================================
    add_heading(doc, "1. Introduction", level=1)

    add_heading(doc, "1.1 Purpose of the System", level=2)
    add_body(doc,
        "The Intelligent Controls & Automation (ICA) Service Management Information System "
        "is designed to digitize and streamline the internal operational workflows of ICA — "
        "a Karachi-based industrial automation and controls engineering firm. The system "
        "provides a centralized, web-based platform for managing engineering parts inventory, "
        "corporate client records, field service orders, preventive maintenance schedules, "
        "and billing invoices."
    )

    add_heading(doc, "1.2 Problem Being Solved", level=2)
    add_body(doc,
        "Prior to this system, ICA relied on manual spreadsheets and paper-based records "
        "to track inventory levels, client accounts, and scheduled maintenance. This led to:"
    )
    add_bullet(doc, "Inaccurate or outdated stock counts for critical parts (PLCs, Sensors, Motor Drives).")
    add_bullet(doc, "Missed preventive maintenance appointments due to lack of scheduling visibility.")
    add_bullet(doc, "Difficulty in generating invoices for completed service orders.")
    add_bullet(doc, "No audit trail for record changes, making accountability difficult.")
    add_body(doc, "The system eliminates these problems by providing a single source of truth backed by a relational PostgreSQL database.")

    add_heading(doc, "1.3 Scope of the Project", level=2)
    add_body(doc, "The system covers the following functional areas:")
    items = [
        "Parts Inventory Management (PLCs, Sensors, Transistors, Diodes, Motor Drives)",
        "Corporate Client Directory",
        "Engineering Service Order Tracking",
        "Preventive Maintenance Scheduling",
        "Invoice Generation for completed service orders",
        "Supplier Management",
        "Audit Logging for all write operations",
        "Role-based access control (Admin vs. Normal User)",
    ]
    for item in items:
        add_bullet(doc, item)

    add_heading(doc, "1.4 Company Whose Problem Is Being Solved", level=2)
    add_body(doc,
        "Company: Intelligent Controls & Automation (ICA)\n"
        "Industry: Industrial Automation & Controls Engineering\n"
        "Location: Plot 12, Block C, Korangi Industrial Area, Karachi, Pakistan\n"
        "Contact: billing@ica-pakistan.com | +92-21-3456-7890\n"
        "NTN: NTN-1234567-8"
    )
    add_body(doc,
        "ICA provides PLC programming, sensor integration, motor drive installation, and "
        "preventive maintenance services to major industrial clients in Karachi, including "
        "Karachi Textile Mills, Pak-Suzuki, Steel Authority of Pakistan, and others."
    )

    doc.add_page_break()

    # =========================================================================
    # 2. SYSTEM OVERVIEW
    # =========================================================================
    add_heading(doc, "2. System Overview", level=1)

    add_heading(doc, "2.1 Application Description", level=2)
    add_body(doc,
        "The ICA Service Management System is a full-stack web application that provides "
        "a dark-themed, responsive dashboard interface for ICA staff. It allows administrators "
        "and normal users to interact with the company's operational database in real time. "
        "The application is split into a React.js Single Page Application (SPA) frontend and "
        "a Flask REST API backend, connected to a cloud-hosted PostgreSQL database on Supabase."
    )

    add_heading(doc, "2.2 Technology Stack", level=2)
    headers = ["Layer", "Technology", "Purpose"]
    rows = [
        ["Database", "PostgreSQL 15 (hosted on Supabase)", "Relational data storage, constraints, views"],
        ["ORM / DB Client", "supabase-py (Python SDK)", "Database interaction from Flask"],
        ["Backend", "Python 3.11 + Flask 3.x", "REST API, business logic, validation"],
        ["CORS", "flask-cors", "Cross-origin request handling"],
        ["Frontend", "React.js 18 (Vite)", "Single Page Application"],
        ["UI Icons", "lucide-react", "Icon library for UI components"],
        ["HTTP Client (FE)", "Axios", "API calls from React to Flask"],
        ["Styling", "Vanilla CSS + custom design system", "Dark-mode glassmorphism UI"],
        ["Hosting (DB)", "Supabase Cloud (us-east-1)", "Managed PostgreSQL with Auth & RLS"],
    ]
    add_schema_table(doc, headers, rows)

    add_heading(doc, "2.3 User Roles and Access", level=2)
    add_body(doc, "The system supports two distinct user roles:")

    add_body(doc, "Administrator (Admin):", bold=True)
    admin_perms = [
        "Full CRUD access: Create, Read, Update, Delete on Inventory, Clients, Service Orders, and Maintenance Schedules",
        "Can change order status (Pending → In Progress → Completed → Cancelled)",
        "Can add and delete inventory items (including hard-delete with audit log)",
        "Can add new corporate clients and update their contact information",
        "Can generate invoices for completed service orders",
        "Can view the full audit log",
        "Export inventory data as CSV",
    ]
    for p in admin_perms:
        add_bullet(doc, p)

    add_body(doc, "Normal User (Read + Limited Write):", bold=True)
    user_perms = [
        "Read access to all modules: Inventory, Clients, Orders, Maintenance, Suppliers",
        "Can search and filter inventory by part name or National Stock Number (NSN)",
        "Can filter service orders by status",
        "Can export inventory to CSV",
        "Cannot add, edit, or delete any records",
        "Action buttons (Add, Edit, Delete) are hidden from the UI",
    ]
    for p in user_perms:
        add_bullet(doc, p)

    add_heading(doc, "2.4 External Views Per User", level=2)
    add_body(doc, "Admin View:", bold=True)
    add_bullet(doc, "Parts Inventory table with Edit (pencil icon) and Delete (trash icon) buttons per row")
    add_bullet(doc, "'Add Item' button to insert new parts")
    add_bullet(doc, "Service Orders with status dropdown to change order state")
    add_bullet(doc, "Corporate Clients with 'Add Client' and 'Edit' buttons per card")
    add_bullet(doc, "Maintenance Schedules with 'Add Schedule' button")
    add_bullet(doc, "Invoice generation button for completed orders")
    add_bullet(doc, "Audit Log page showing who changed what and when")
    doc.add_paragraph()
    add_body(doc, "Normal User View:", bold=True)
    add_bullet(doc, "Read-only inventory table with search and category filter tabs")
    add_bullet(doc, "Client directory (card grid) — view only")
    add_bullet(doc, "Service orders list filtered by status — view only")
    add_bullet(doc, "Maintenance schedules calendar/list — view only")
    add_bullet(doc, "CSV export of inventory visible to all users")

    doc.add_page_break()

    # =========================================================================
    # 3. DATABASE DESIGN
    # =========================================================================
    add_heading(doc, "3. Database Design", level=1)
    add_body(doc,
        "The database is hosted on Supabase (managed PostgreSQL 15). It consists of seven "
        "tables designed using normalization principles to minimize redundancy and ensure "
        "referential integrity. The entity-relationship structure models the real-world "
        "operations of ICA."
    )

    add_heading(doc, "3.1 Tables & Schema", level=2)

    # --- inventory
    add_body(doc, "Table: inventory", bold=True)
    add_body(doc, "Stores all engineering parts held in stock by ICA.")
    headers = ["Column", "Data Type", "Constraint", "Description"]
    rows = [
        ["item_id", "int8", "PRIMARY KEY, IDENTITY", "Unique identifier for each inventory item"],
        ["name", "text", "NOT NULL", "Name of the part (e.g., Siemens S7-1200 PLC)"],
        ["category", "text", "NOT NULL, CHECK", "Category: PLC, Sensor, Transistor, Diode, Motor Drive"],
        ["supplier_id", "int8", "NULLABLE, FK → suppliers", "Reference to the supplying company"],
        ["quantity", "int8", "NOT NULL, CHECK ≥ 0", "Current stock quantity"],
        ["unit_price", "float8", "NOT NULL, CHECK > 0", "Price per unit in PKR"],
        ["nsn", "text", "NULLABLE", "NATO National Stock Number for identification"],
    ]
    add_schema_table(doc, headers, rows)

    # --- suppliers
    add_body(doc, "Table: suppliers", bold=True)
    add_body(doc, "Stores supplier/vendor information for parts procurement.")
    rows = [
        ["supplier_id", "int8", "PRIMARY KEY, IDENTITY", "Unique supplier identifier"],
        ["name", "text", "NOT NULL", "Supplier company name"],
        ["contact", "text", "NULLABLE", "Contact details of the supplier"],
        ["category", "text", "NULLABLE", "Type of parts supplied"],
    ]
    add_schema_table(doc, headers, rows)

    # --- Clients
    add_body(doc, "Table: Clients", bold=True)
    add_body(doc, "Stores corporate clients who receive engineering services from ICA.")
    rows = [
        ["client_id", "int8", "PRIMARY KEY, IDENTITY", "Unique client identifier"],
        ["company_name", "text", "NOT NULL", "Name of the client's company"],
        ["contact_person", "text", "NULLABLE", "Name of the primary contact at the client side"],
        ["email", "text", "NULLABLE", "Email address of the contact person"],
    ]
    add_schema_table(doc, headers, rows)

    # --- Service_Orders
    add_body(doc, "Table: Service_Orders", bold=True)
    add_body(doc, "Tracks engineering service orders raised by clients.")
    rows = [
        ["order_id", "int8", "PRIMARY KEY, IDENTITY", "Unique order identifier"],
        ["client_id", "int8", "NOT NULL, FK → Clients", "References which client raised this order"],
        ["description", "text", "NOT NULL", "Description of the engineering work to be done"],
        ["order_date", "text", "NOT NULL", "Date the order was placed (YYYY-MM-DD)"],
        ["status", "text", "NOT NULL, CHECK", "Status: Pending, In Progress, Completed, Cancelled"],
        ["total_cost", "float8", "NOT NULL, CHECK ≥ 0", "Total service cost in PKR"],
    ]
    add_schema_table(doc, headers, rows)

    # --- Invoices
    add_body(doc, "Table: Invoices", bold=True)
    add_body(doc, "Stores billing invoice records tied to completed service orders.")
    rows = [
        ["invoice_id", "int8", "PRIMARY KEY, IDENTITY", "Unique invoice identifier"],
        ["order_id", "int8", "NOT NULL, FK → Service_Orders", "References the completed service order"],
        ["invoice_date", "date", "NOT NULL", "Date the invoice was issued"],
        ["amount_due", "float8", "NOT NULL", "Total amount due (may include tax)"],
        ["payment_status", "text", "NULLABLE", "Status: Paid, Unpaid, Overdue"],
    ]
    add_schema_table(doc, headers, rows)

    # --- Maintenance_Schedules
    add_body(doc, "Table: Maintenance_Schedules", bold=True)
    add_body(doc, "Tracks preventive maintenance schedules for client equipment.")
    rows = [
        ["schedule_id", "int8", "PRIMARY KEY, IDENTITY", "Unique schedule identifier"],
        ["client_id", "int8", "NOT NULL, FK → Clients", "Client whose equipment is being maintained"],
        ["equipment_name", "text", "NOT NULL", "Name/model of the equipment"],
        ["last_service_date", "date", "NULLABLE", "Date of last maintenance performed"],
        ["next_service_date", "date", "NOT NULL", "Date of next scheduled maintenance"],
        ["technician_assigned", "text", "NULLABLE", "Name of the ICA technician assigned"],
    ]
    add_schema_table(doc, headers, rows)

    # --- audit_logs
    add_body(doc, "Table: audit_logs", bold=True)
    add_body(doc, "Automatically records every INSERT, UPDATE, and DELETE operation for compliance and accountability.")
    rows = [
        ["id", "int8", "PRIMARY KEY, IDENTITY", "Unique log entry identifier"],
        ["table_name", "text", "NOT NULL", "Name of the table that was modified"],
        ["record_id", "text", "NOT NULL", "Primary key of the modified record"],
        ["action", "text", "NOT NULL", "Action performed: INSERT, UPDATE, DELETE"],
        ["changed_by", "text", "NULLABLE", "Role or user identifier who performed the action"],
        ["changed_at", "timestamptz", "NOT NULL, DEFAULT now()", "Timestamp of the action (with timezone)"],
    ]
    add_schema_table(doc, headers, rows)

    add_heading(doc, "3.2 Constraints", level=2)
    add_body(doc, "The following integrity constraints are enforced at the database level:")

    constraints = [
        ("PRIMARY KEY", "Every table has an integer primary key with IDENTITY (auto-increment). This guarantees uniqueness and serves as the basis for foreign key references."),
        ("FOREIGN KEY", "• inventory.supplier_id → suppliers.supplier_id\n• Service_Orders.client_id → Clients.client_id\n• Maintenance_Schedules.client_id → Clients.client_id\n• Invoices.order_id → Service_Orders.order_id\nThese enforce referential integrity — you cannot add an order for a non-existent client, or create an invoice for a non-existent order."),
        ("NOT NULL", "Critical fields such as name, category, quantity, unit_price (inventory), description, status, total_cost (orders), equipment_name, next_service_date (maintenance), company_name (clients), and all audit_logs fields are marked NOT NULL to prevent incomplete records."),
        ("CHECK (category)", "The inventory.category column is validated at the application layer (Flask) to only accept: PLC, Sensor, Transistor, Diode, Motor Drive."),
        ("CHECK (status)", "Service_Orders.status is validated to only accept: Pending, In Progress, Completed, Cancelled."),
        ("CHECK (quantity ≥ 0)", "Inventory quantity cannot go below zero, enforced in Flask before database writes."),
        ("CHECK (unit_price > 0)", "Unit price must be strictly positive, enforced in Flask before writes."),
        ("Email Format Validation", "The client's email address is validated in Flask using a basic format check (must contain '@' and '.') before being written to the database."),
        ("Read-only fields", "Primary keys, NSN, and part names are protected from update in the API layer — attempts to modify them are rejected with HTTP 400."),
        ("Audit Logging Constraint", "Every INSERT, UPDATE, and DELETE operation triggers a corresponding entry in audit_logs, recording the table, record ID, action type, and the role of the user who performed it."),
    ]

    for constraint_name, description in constraints:
        add_body(doc, f"{constraint_name}:", bold=True)
        add_body(doc, description)
        doc.add_paragraph()

    doc.add_page_break()

    # =========================================================================
    # 4. DATA MANIPULATION OPERATIONS
    # =========================================================================
    add_heading(doc, "4. Data Manipulation Operations", level=1)
    add_body(doc,
        "All data manipulation is performed through the Flask REST API, which validates inputs "
        "before executing Supabase PostgreSQL queries. The React frontend sends HTTP requests "
        "with JSON payloads to the API endpoints."
    )

    add_heading(doc, "4.1 Insert (CREATE)", level=2)
    add_body(doc, "New records can be inserted across three primary modules:")

    add_body(doc, "a) Adding a New Inventory Item:", bold=True)
    add_body(doc, "Admin clicks 'Add Item' button → modal form opens → fills in Part Name, Category, NSN, Quantity, and Unit Price → submits → Flask validates → inserts into inventory table.")
    add_code_block(doc, "POST /api/inventory")
    add_code_block(doc, 'Body: {"name":"Siemens S7-1200 PLC","category":"PLC","quantity":10,"unit_price":45000,"nsn":"5961-00-123-4567"}')
    add_code_block(doc, "SQL: INSERT INTO inventory (name, category, quantity, unit_price, nsn) VALUES (...) RETURNING *;")
    doc.add_paragraph()

    add_body(doc, "b) Adding a New Service Order:", bold=True)
    add_body(doc, "Admin opens 'Add Order' → fills description, client, cost, status → submits → Flask validates client_id exists → inserts.")
    add_code_block(doc, "POST /api/orders")
    add_code_block(doc, 'Body: {"description":"PLC Installation at Pak-Suzuki","client_id":3,"total_cost":85000,"status":"Pending"}')
    add_code_block(doc, "SQL: INSERT INTO Service_Orders (description, client_id, total_cost, status, order_date) VALUES (...);")
    doc.add_paragraph()

    add_body(doc, "c) Adding a New Corporate Client:", bold=True)
    add_code_block(doc, "POST /api/clients")
    add_code_block(doc, 'Body: {"company_name":"Fauji Fertilizer Co.","contact_person":"Bilal Khan","email":"bilal@fauji.com"}')
    add_code_block(doc, "SQL: INSERT INTO Clients (company_name, contact_person, email) VALUES (...);")
    doc.add_paragraph()

    add_body(doc, "d) Adding a Maintenance Schedule:", bold=True)
    add_code_block(doc, "POST /api/maintenance")
    add_code_block(doc, 'Body: {"equipment_name":"Siemens SINAMICS Drive","client_id":2,"technician_assigned":"Ahmed Raza","next_service_date":"2026-07-15"}')
    add_code_block(doc, "SQL: INSERT INTO Maintenance_Schedules (...) VALUES (...);")
    doc.add_paragraph()
    add_body(doc, "All insertions are followed by an automatic audit log entry recording the action, affected table, and user role.")

    add_heading(doc, "4.2 Update (UPDATE)", level=2)
    add_body(doc, "Updates are restricted to non-key fields to protect data integrity:")

    add_body(doc, "a) Updating Inventory Item:", bold=True)
    add_body(doc, "Admin clicks 'Edit' (pencil icon) on a row → modal shows locked fields (ID, Name, NSN) → can update Quantity, Unit Price, Category → saves.")
    add_code_block(doc, "PUT /api/inventory/<item_id>")
    add_code_block(doc, 'Body: {"quantity": 25, "unit_price": 47500.00, "category": "PLC"}')
    add_code_block(doc, "SQL: UPDATE inventory SET quantity=25, unit_price=47500, category='PLC' WHERE item_id=<id>;")
    doc.add_paragraph()

    add_body(doc, "b) Updating Order Status:", bold=True)
    add_body(doc, "Admin selects a new status from the dropdown → only the status field is updated.")
    add_code_block(doc, "PUT /api/orders/<order_id>")
    add_code_block(doc, 'Body: {"status": "Completed"}')
    add_code_block(doc, "SQL: UPDATE Service_Orders SET status='Completed' WHERE order_id=<id>;")
    doc.add_paragraph()

    add_body(doc, "c) Updating Client Contact Information:", bold=True)
    add_code_block(doc, "PUT /api/clients/<client_id>")
    add_code_block(doc, 'Body: {"contact_person": "Ali Nawaz", "email": "ali@paksuzu.com"}')
    add_code_block(doc, "SQL: UPDATE Clients SET contact_person='Ali Nawaz', email='ali@paksuzu.com' WHERE client_id=<id>;")

    add_heading(doc, "4.3 Delete (DELETE)", level=2)
    add_body(doc,
        "Delete is restricted to Admin users only. This is enforced both at the frontend "
        "(button is hidden for non-admins) and at the backend (HTTP 403 Forbidden is returned "
        "if the X-User-Role header is not 'admin')."
    )
    add_body(doc, "Deleting an Inventory Item:", bold=True)
    add_body(doc, "Admin clicks the 'Trash' icon on an inventory row → confirmation dialog appears → confirms → item is deleted.")
    add_code_block(doc, "DELETE /api/inventory/<item_id>")
    add_code_block(doc, "SQL: DELETE FROM inventory WHERE item_id=<id>;")
    add_body(doc, "An audit log entry is always created after a successful delete.")

    add_heading(doc, "4.4 Search & Filter", level=2)
    add_body(doc, "The system provides powerful, real-time search and filtering capabilities:")

    add_body(doc, "a) Inventory Search (by Part Name or NSN):", bold=True)
    add_code_block(doc, "GET /api/inventory/search?q=siemens")
    add_code_block(doc, "SQL: SELECT * FROM inventory WHERE name ILIKE '%siemens%' OR nsn ILIKE '%siemens%';")
    add_body(doc, "Search is debounced (350ms) in the frontend to reduce unnecessary API calls. Results from both name and NSN queries are merged and de-duplicated.")
    doc.add_paragraph()

    add_body(doc, "b) Inventory Category Filter:", bold=True)
    add_code_block(doc, "GET /api/inventory?category=PLC")
    add_code_block(doc, "SQL: SELECT * FROM inventory WHERE category='PLC' ORDER BY item_id;")
    doc.add_paragraph()

    add_body(doc, "c) Service Orders Status Filter:", bold=True)
    add_code_block(doc, "GET /api/orders?status=Pending")
    add_code_block(doc, "SQL: SELECT * FROM Service_Orders WHERE status='Pending' ORDER BY order_date DESC;")
    doc.add_paragraph()

    add_body(doc, "d) Client Search (Frontend-side):", bold=True)
    add_body(doc, "All clients are fetched at once and filtered in the browser by company name or contact person using JavaScript's .filter() method — no additional API call needed.")

    doc.add_page_break()

    # =========================================================================
    # 5. QUERY IMPLEMENTATION
    # =========================================================================
    add_heading(doc, "5. Query Implementation", level=1)
    add_body(doc, "The following SQL queries form the core of the system's data retrieval and reporting logic.")

    add_heading(doc, "5.1 Basic SELECT Queries", level=2)

    add_body(doc, "Q1 — Fetch all inventory items ordered by ID:", bold=True)
    add_code_block(doc, "SELECT item_id, name, category, quantity, unit_price, nsn, supplier_id")
    add_code_block(doc, "FROM inventory")
    add_code_block(doc, "ORDER BY item_id;")
    doc.add_paragraph()

    add_body(doc, "Q2 — Fetch all clients:", bold=True)
    add_code_block(doc, "SELECT client_id, company_name, contact_person, email")
    add_code_block(doc, "FROM Clients")
    add_code_block(doc, "ORDER BY client_id;")
    doc.add_paragraph()

    add_body(doc, "Q3 — Fetch all suppliers:", bold=True)
    add_code_block(doc, "SELECT * FROM suppliers ORDER BY supplier_id;")
    doc.add_paragraph()

    add_body(doc, "Q4 — Fetch all maintenance schedules ordered by next service date:", bold=True)
    add_code_block(doc, "SELECT schedule_id, equipment_name, last_service_date,")
    add_code_block(doc, "       next_service_date, technician_assigned, client_id")
    add_code_block(doc, "FROM Maintenance_Schedules")
    add_code_block(doc, "ORDER BY next_service_date ASC;")
    doc.add_paragraph()

    add_heading(doc, "5.2 JOIN Queries", level=2)

    add_body(doc, "Q5 — Fetch all service orders with client company name (INNER JOIN):", bold=True)
    add_code_block(doc, "SELECT so.order_id, so.description, so.order_date, so.status,")
    add_code_block(doc, "       so.total_cost, c.company_name")
    add_code_block(doc, "FROM Service_Orders so")
    add_code_block(doc, "INNER JOIN Clients c ON so.client_id = c.client_id")
    add_code_block(doc, "ORDER BY so.order_date DESC;")
    doc.add_paragraph()

    add_body(doc, "Q6 — Fetch maintenance schedules with client names (INNER JOIN):", bold=True)
    add_code_block(doc, "SELECT ms.schedule_id, ms.equipment_name, ms.last_service_date,")
    add_code_block(doc, "       ms.next_service_date, ms.technician_assigned, c.company_name")
    add_code_block(doc, "FROM Maintenance_Schedules ms")
    add_code_block(doc, "INNER JOIN Clients c ON ms.client_id = c.client_id")
    add_code_block(doc, "ORDER BY ms.next_service_date ASC;")
    doc.add_paragraph()

    add_body(doc, "Q7 — Fetch invoice details with order and client info (Multi-table JOIN):", bold=True)
    add_code_block(doc, "SELECT i.invoice_id, i.invoice_date, i.amount_due, i.payment_status,")
    add_code_block(doc, "       so.description AS order_description, so.order_date, so.total_cost,")
    add_code_block(doc, "       c.company_name, c.contact_person, c.email")
    add_code_block(doc, "FROM Invoices i")
    add_code_block(doc, "INNER JOIN Service_Orders so ON i.order_id = so.order_id")
    add_code_block(doc, "INNER JOIN Clients c ON so.client_id = c.client_id")
    add_code_block(doc, "ORDER BY i.invoice_date DESC;")
    doc.add_paragraph()

    add_body(doc, "Q8 — Inventory with supplier name (LEFT JOIN to include unsupplied items):", bold=True)
    add_code_block(doc, "SELECT inv.item_id, inv.name, inv.category, inv.quantity,")
    add_code_block(doc, "       inv.unit_price, inv.nsn, s.name AS supplier_name")
    add_code_block(doc, "FROM inventory inv")
    add_code_block(doc, "LEFT JOIN suppliers s ON inv.supplier_id = s.supplier_id")
    add_code_block(doc, "ORDER BY inv.item_id;")
    doc.add_paragraph()

    add_heading(doc, "5.3 Filtering and Sorting Queries", level=2)

    add_body(doc, "Q9 — Filter inventory by category:", bold=True)
    add_code_block(doc, "SELECT * FROM inventory WHERE category = 'PLC' ORDER BY item_id;")
    doc.add_paragraph()

    add_body(doc, "Q10 — Search inventory by part name or NSN (case-insensitive):", bold=True)
    add_code_block(doc, "SELECT item_id, name, category, quantity, unit_price, nsn")
    add_code_block(doc, "FROM inventory")
    add_code_block(doc, "WHERE name ILIKE '%siemens%' OR nsn ILIKE '%5961%';")
    doc.add_paragraph()

    add_body(doc, "Q11 — Filter service orders by status:", bold=True)
    add_code_block(doc, "SELECT * FROM Service_Orders WHERE status = 'In Progress' ORDER BY order_date DESC;")
    doc.add_paragraph()

    add_body(doc, "Q12 — Identify low-stock items (quantity < 10 but > 0):", bold=True)
    add_code_block(doc, "SELECT item_id, name, category, quantity")
    add_code_block(doc, "FROM inventory")
    add_code_block(doc, "WHERE quantity > 0 AND quantity < 10")
    add_code_block(doc, "ORDER BY quantity ASC;")
    doc.add_paragraph()

    add_body(doc, "Q13 — Identify out-of-stock items:", bold=True)
    add_code_block(doc, "SELECT item_id, name, category FROM inventory WHERE quantity = 0;")
    doc.add_paragraph()

    add_heading(doc, "5.4 Aggregate / Report Queries", level=2)

    add_body(doc, "Q14 — Total inventory value per category:", bold=True)
    add_code_block(doc, "SELECT category,")
    add_code_block(doc, "       COUNT(*) AS total_items,")
    add_code_block(doc, "       SUM(quantity) AS total_units,")
    add_code_block(doc, "       SUM(quantity * unit_price) AS total_value_pkr")
    add_code_block(doc, "FROM inventory")
    add_code_block(doc, "GROUP BY category")
    add_code_block(doc, "ORDER BY total_value_pkr DESC;")
    doc.add_paragraph()

    add_body(doc, "Q15 — Count of service orders per client:", bold=True)
    add_code_block(doc, "SELECT c.company_name, COUNT(so.order_id) AS total_orders,")
    add_code_block(doc, "       SUM(so.total_cost) AS total_revenue")
    add_code_block(doc, "FROM Clients c")
    add_code_block(doc, "LEFT JOIN Service_Orders so ON c.client_id = so.client_id")
    add_code_block(doc, "GROUP BY c.company_name")
    add_code_block(doc, "ORDER BY total_revenue DESC;")
    doc.add_paragraph()

    add_body(doc, "Q16 — Upcoming maintenance in the next 30 days:", bold=True)
    add_code_block(doc, "SELECT ms.equipment_name, ms.next_service_date,")
    add_code_block(doc, "       ms.technician_assigned, c.company_name")
    add_code_block(doc, "FROM Maintenance_Schedules ms")
    add_code_block(doc, "INNER JOIN Clients c ON ms.client_id = c.client_id")
    add_code_block(doc, "WHERE ms.next_service_date BETWEEN CURRENT_DATE AND CURRENT_DATE + INTERVAL '30 days'")
    add_code_block(doc, "ORDER BY ms.next_service_date;")
    doc.add_paragraph()

    add_body(doc, "Q17 — Audit log for a specific table:", bold=True)
    add_code_block(doc, "SELECT id, record_id, action, changed_by, changed_at")
    add_code_block(doc, "FROM audit_logs")
    add_code_block(doc, "WHERE table_name = 'inventory'")
    add_code_block(doc, "ORDER BY changed_at DESC;")
    doc.add_paragraph()

    add_heading(doc, "5.5 Database Views", level=2)
    add_body(doc,
        "The following view is recommended for simplified invoice reporting (can be created in Supabase SQL Editor):"
    )
    add_code_block(doc, "CREATE VIEW vw_invoice_summary AS")
    add_code_block(doc, "SELECT i.invoice_id,")
    add_code_block(doc, "       CONCAT('ICA-INV-', LPAD(i.order_id::text, 4, '0')) AS invoice_number,")
    add_code_block(doc, "       c.company_name, c.email,")
    add_code_block(doc, "       so.description AS service_description,")
    add_code_block(doc, "       so.total_cost AS subtotal,")
    add_code_block(doc, "       ROUND(so.total_cost * 0.17, 2) AS gst_17pct,")
    add_code_block(doc, "       ROUND(so.total_cost * 1.17, 2) AS grand_total,")
    add_code_block(doc, "       i.payment_status, i.invoice_date")
    add_code_block(doc, "FROM Invoices i")
    add_code_block(doc, "JOIN Service_Orders so ON i.order_id = so.order_id")
    add_code_block(doc, "JOIN Clients c ON so.client_id = c.client_id;")

    doc.add_page_break()

    # =========================================================================
    # 6. FRONTEND & DATABASE CONNECTIVITY
    # =========================================================================
    add_heading(doc, "6. Frontend & Database Connectivity", level=1)

    add_heading(doc, "6.1 Frontend Interface Description", level=2)
    add_body(doc,
        "The frontend is a React.js Single Page Application (SPA) built with Vite. It uses a "
        "dark-mode glassmorphism design with a sidebar navigation and content panels. The following "
        "main screens are implemented:"
    )
    screens = [
        ("Parts Inventory", "Data table with search bar (NSN/name), category filter tabs (PLC, Sensor, Transistor, Diode, Motor Drive), stock status badges, and admin action buttons (Edit, Delete, Add Item). CSV export available to all users."),
        ("Corporate Clients", "Card-grid layout showing company name, contact person, email. Admin can add new clients or edit contact info. Users can search by company or contact."),
        ("Service Orders", "Table showing all engineering service orders with client name (joined via FK), order date, status badge, and cost. Admin can change status via dropdown."),
        ("Maintenance Schedules", "List of upcoming and past maintenance appointments, sorted by next_service_date. Shows assigned technician and client."),
        ("Invoice Generation", "Admin selects a completed order → clicks Generate Invoice → system fetches order + client + invoice data and renders a formatted invoice (invoice number, company details, amounts, 17% GST)."),
        ("Suppliers", "Simple read-only list of suppliers with name, contact, and category."),
    ]
    for name, desc in screens:
        add_body(doc, f"{name}:", bold=True)
        add_body(doc, desc)

    add_heading(doc, "6.2 Frontend-to-Database Connection Flow", level=2)
    add_body(doc, "The connection architecture is as follows:")
    add_bullet(doc, "Step 1: React component calls a function in api.js (Axios-based service layer).")
    add_bullet(doc, "Step 2: Axios sends an HTTP request (GET/POST/PUT/DELETE) to the Flask REST API.")
    add_bullet(doc, "Step 3: Flask validates the request body, checks constraints, and builds a Supabase query.")
    add_bullet(doc, "Step 4: supabase-py SDK translates the query into a PostgreSQL REST call to Supabase.")
    add_bullet(doc, "Step 5: Supabase executes SQL on the managed PostgreSQL instance and returns JSON.")
    add_bullet(doc, "Step 6: Flask returns the JSON response with appropriate HTTP status code.")
    add_bullet(doc, "Step 7: React state is updated and the UI re-renders to reflect the new data.")
    doc.add_paragraph()
    add_body(doc, "The X-User-Role HTTP header is sent with every write request from the frontend to identify the user's role to the Flask API, which uses it for permission checks and audit logging.")

    add_heading(doc, "6.3 Admin Panel vs. Normal User Interface", level=2)
    add_body(doc, "The React component reads the user.role property from the login state and conditionally renders action buttons:")
    add_code_block(doc, "const isAdmin = user?.role === 'admin'")
    add_code_block(doc, "{isAdmin && <button>Add Item</button>}   // Only shown to admin")
    add_code_block(doc, "{isAdmin && <button>Delete</button>}     // Only shown to admin")
    add_body(doc,
        "This means the Admin Panel and Normal User Panel are the same React application — "
        "the difference is entirely in what UI elements are rendered based on the logged-in user's role."
    )

    doc.add_page_break()

    # =========================================================================
    # 7. REPORTS GENERATED
    # =========================================================================
    add_heading(doc, "7. Reports Generated", level=1)
    add_body(doc, "The system generates the following reports and data exports:")

    reports = [
        (
            "Invoice Report (PDF-printable JSON)",
            "Generated via POST /api/generate-invoice. Includes the invoice number (ICA-INV-XXXX), company details, client information, service description, subtotal, 17% GST, and grand total. Only available for Completed service orders. Accessible by Admin.",
            "Admin only"
        ),
        (
            "Inventory CSV Export",
            "All inventory items (item_id, name, category, NSN, quantity, unit_price) exported as a CSV file with a timestamped filename. Available to all users via the 'Export CSV' button.",
            "All users"
        ),
        (
            "Low-Stock Alert Report",
            "Real-time visual report displayed in the Inventory module. Items with quantity between 1 and 9 are highlighted orange ('Low Stock'). Items with quantity = 0 are highlighted red ('Out of Stock').",
            "All users"
        ),
        (
            "Service Order Status Report",
            "The Orders module displays all service orders filterable by status (Pending, In Progress, Completed, Cancelled). Shows totals and client names via JOIN.",
            "All users"
        ),
        (
            "Maintenance Schedule Report",
            "The Maintenance module lists all upcoming maintenance appointments sorted by next_service_date. Shows the assigned technician and client company for each scheduled event.",
            "All users"
        ),
        (
            "Audit Trail Report",
            "The audit_logs table records every INSERT, UPDATE, and DELETE with timestamp, role, table name, and record ID. This report is accessible via the Supabase dashboard and can be queried via SQL.",
            "Admin / DBA"
        ),
    ]

    headers = ["Report", "Description", "Access"]
    rows_r = [[name, desc, access] for name, desc, access in reports]
    add_schema_table(doc, headers, rows_r)

    doc.add_page_break()

    # =========================================================================
    # 8. USER ROLES & PERMISSIONS
    # =========================================================================
    add_heading(doc, "8. User Roles & Permissions", level=1)
    add_body(doc, "The system implements a two-role access control model:")

    headers = ["Feature / Action", "Admin", "Normal User"]
    rows = [
        ["View Inventory", "✓", "✓"],
        ["Search Inventory (by Name/NSN)", "✓", "✓"],
        ["Filter Inventory by Category", "✓", "✓"],
        ["Export Inventory CSV", "✓", "✓"],
        ["Add New Inventory Item", "✓", "✗"],
        ["Edit Inventory Item (Qty/Price)", "✓", "✗"],
        ["Delete Inventory Item", "✓", "✗"],
        ["View Corporate Clients", "✓", "✓"],
        ["Search Clients", "✓", "✓"],
        ["Add New Client", "✓", "✗"],
        ["Edit Client Contact Info", "✓", "✗"],
        ["View Service Orders", "✓", "✓"],
        ["Filter Orders by Status", "✓", "✓"],
        ["Create New Order", "✓", "✗"],
        ["Update Order Status", "✓", "✗"],
        ["View Maintenance Schedules", "✓", "✓"],
        ["Add Maintenance Schedule", "✓", "✗"],
        ["Generate Invoice", "✓", "✗"],
        ["View Suppliers", "✓", "✓"],
        ["View Audit Log", "✓ (via Supabase)", "✗"],
    ]
    add_schema_table(doc, headers, rows)

    add_body(doc,
        "Role enforcement is implemented at two layers: (1) Frontend — action buttons are "
        "conditionally rendered based on user.role. (2) Backend — the X-User-Role header is "
        "checked for delete operations, returning HTTP 403 Forbidden if the role is not 'admin'."
    )

    doc.add_page_break()

    # =========================================================================
    # 9. TESTING
    # =========================================================================
    add_heading(doc, "9. Testing", level=1)

    add_heading(doc, "9.1 Testing Approach", level=2)
    add_body(doc,
        "The system was tested through manual functional testing using the browser interface "
        "and direct API testing with REST client tools. Testing covered all CRUD operations, "
        "role-based access checks, input validation, and error handling."
    )

    add_heading(doc, "9.2 Test Cases", level=2)
    headers = ["Test ID", "Test Case", "Input / Action", "Expected Output", "Result"]
    rows = [
        ["TC-01", "Insert new inventory item", "POST /api/inventory with valid fields", "201 Created, item appears in list", "Pass"],
        ["TC-02", "Insert with missing required field", "POST /api/inventory without 'name'", "400 Bad Request with error message", "Pass"],
        ["TC-03", "Insert with invalid category", "category='Unknown'", "400 Bad Request: invalid category", "Pass"],
        ["TC-04", "Update inventory quantity", "PUT /api/inventory/1 {quantity:50}", "200 OK, quantity updated in DB", "Pass"],
        ["TC-05", "Update with negative quantity", "PUT with quantity=-5", "400 Bad Request", "Pass"],
        ["TC-06", "Delete inventory item (Admin)", "DELETE /api/inventory/5 with role=admin", "200 OK, item removed, audit log updated", "Pass"],
        ["TC-07", "Delete by non-admin", "DELETE with role=user", "403 Forbidden", "Pass"],
        ["TC-08", "Search by part name", "GET /api/inventory/search?q=siemens", "Returns matching items (case-insensitive)", "Pass"],
        ["TC-09", "Search by NSN", "GET /api/inventory/search?q=5961", "Returns items with matching NSN", "Pass"],
        ["TC-10", "Filter by category", "GET /api/inventory?category=PLC", "Only PLC items returned", "Pass"],
        ["TC-11", "Filter orders by status", "GET /api/orders?status=Pending", "Only pending orders returned", "Pass"],
        ["TC-12", "Create service order", "POST /api/orders with valid body", "201 Created, order in DB", "Pass"],
        ["TC-13", "Update order status", "PUT /api/orders/1 {status:Completed}", "200 OK, status updated", "Pass"],
        ["TC-14", "Invalid order status", "PUT with status='Done'", "400 Bad Request", "Pass"],
        ["TC-15", "Generate invoice (completed order)", "POST /api/generate-invoice {order_id:3}", "200 OK, invoice JSON returned", "Pass"],
        ["TC-16", "Generate invoice (non-completed order)", "POST with order_id for Pending order", "400 Bad Request: order not completed", "Pass"],
        ["TC-17", "Add new client", "POST /api/clients with valid fields", "201 Created", "Pass"],
        ["TC-18", "Invalid email format", "POST /api/clients {email:'invalid'}", "400 Bad Request: invalid email", "Pass"],
        ["TC-19", "Add maintenance schedule", "POST /api/maintenance with valid fields", "201 Created", "Pass"],
        ["TC-20", "Admin UI — Add button visibility", "Login as admin, view Inventory", "Add Item button visible", "Pass"],
        ["TC-21", "User UI — Add button hidden", "Login as user, view Inventory", "Add Item button not rendered", "Pass"],
        ["TC-22", "Low stock visual alert", "Item with qty=5 in DB", "Orange 'Low Stock' badge shown", "Pass"],
        ["TC-23", "Out of stock visual alert", "Item with qty=0 in DB", "Red 'Out of Stock' badge shown", "Pass"],
        ["TC-24", "Audit log entry on insert", "POST to create any record", "New row appears in audit_logs table", "Pass"],
        ["TC-25", "CSV export", "Click Export CSV on Inventory", "CSV file downloaded with timestamp", "Pass"],
    ]
    add_schema_table(doc, headers, rows)

    doc.add_page_break()

    # =========================================================================
    # 10. CONCLUSION
    # =========================================================================
    add_heading(doc, "10. Conclusion", level=1)

    add_heading(doc, "10.1 Summary of Achievements", level=2)
    add_body(doc,
        "This project successfully delivers a fully functional Service Management Information "
        "System for Intelligent Controls & Automation (ICA). The system achieves the following:"
    )
    achievements = [
        "A normalized PostgreSQL database with 7 tables, referential integrity constraints, and audit logging.",
        "A robust Flask REST API with 15+ endpoints covering full CRUD operations with server-side validation.",
        "A responsive, role-aware React.js frontend providing distinct Admin and User experiences.",
        "Real-time inventory search by part name and National Stock Number (NSN).",
        "Automated invoice generation with 17% GST calculation for completed service orders.",
        "Preventive maintenance scheduling with client and technician tracking.",
        "A complete audit trail for all data modifications.",
        "CSV data export for inventory reporting.",
    ]
    for a in achievements:
        add_bullet(doc, a)

    add_heading(doc, "10.2 Challenges Faced", level=2)
    challenges = [
        "Case-sensitive table names in Supabase PostgreSQL (e.g., 'Clients' vs 'clients') required fixing all table references in Flask.",
        "Handling Supabase foreign key joins in the Python SDK required using the nested select syntax (e.g., Clients(company_name)).",
        "Preventing null-value constraint violations when optional fields were passed as null in update payloads — resolved by stripping null values from the patch dictionary.",
        "Implementing real-time search debouncing in React to avoid excessive API calls while typing.",
        "Setting up proper CORS configuration to allow the React dev server to communicate with Flask.",
    ]
    for c in challenges:
        add_bullet(doc, c)

    add_heading(doc, "10.3 Future Improvements", level=2)
    improvements = [
        "Implement Supabase Auth for proper JWT-based authentication and role management instead of the current header-based role system.",
        "Add Row Level Security (RLS) policies in Supabase PostgreSQL to enforce permissions at the database level.",
        "Implement dashboard analytics with charts (e.g., orders per month, revenue by client) using Chart.js or Recharts.",
        "Add email notifications for upcoming maintenance schedules (using Flask-Mail or Supabase Edge Functions).",
        "Implement soft-delete (is_deleted flag) instead of hard-delete for better data recovery.",
        "Add pagination to the inventory and orders tables for scalability with large datasets.",
        "Deploy the Flask backend to a cloud platform (e.g., Render, Railway) and the React frontend to Vercel.",
        "Implement full-text search using PostgreSQL's to_tsvector for faster and more powerful searching.",
    ]
    for imp in improvements:
        add_bullet(doc, imp)

    doc.add_page_break()

    # =========================================================================
    # 11. REFERENCES
    # =========================================================================
    add_heading(doc, "11. References", level=1)
    refs = [
        "Supabase Documentation — https://supabase.com/docs",
        "Flask Documentation (Pallets Projects) — https://flask.palletsprojects.com/",
        "python-docx Library — https://python-docx.readthedocs.io/",
        "React.js Official Documentation — https://react.dev/",
        "PostgreSQL 15 Documentation — https://www.postgresql.org/docs/15/",
        "supabase-py Python SDK — https://github.com/supabase-community/supabase-py",
        "lucide-react Icon Library — https://lucide.dev/",
        "Axios HTTP Client — https://axios-http.com/",
        "Vite Build Tool — https://vitejs.dev/",
        "flask-cors Extension — https://flask-cors.readthedocs.io/",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}] {ref}")
        set_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)

    # =========================================================================
    # Save
    # =========================================================================
    output_path = r"c:\Users\DELL\OneDrive\Desktop\dbms_cep\ICA_DBMS_Final_Report.docx"
    doc.save(output_path)
    print(f"\n✅  Report saved to:\n    {output_path}\n")


if __name__ == "__main__":
    build_report()
